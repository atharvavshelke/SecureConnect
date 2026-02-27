"""
Generate SecureConnect Project Analysis Word Document (Updated MVC Architecture)
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
for level in range(1, 4):
    doc.styles[f'Heading {level}'].font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers), style='Light Shading Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        hdr = table.rows[0].cells[i]
        hdr.text = h
        for p in hdr.paragraphs:
            for r in p.runs: r.bold = True
    for row_data in rows:
        cells = table.add_row().cells
        for i, cell in enumerate(row_data):
            cells[i].text = str(cell)
    doc.add_paragraph()

def mono(text, size=8):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(size)

# ════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════
for _ in range(6): doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('SecureConnect')
r.bold = True; r.font.size = Pt(36); r.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('Project Analysis & Technical Documentation')
r.font.size = Pt(18); r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
doc.add_paragraph()

d = doc.add_paragraph()
d.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = d.add_run('End-to-End Encrypted Chat Application\nMVC Architecture | Sequelize ORM | Socket.IO + Redis')
r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

for _ in range(4): doc.add_paragraph()
i = doc.add_paragraph()
i.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = i.add_run('Node.js | Express.js | Sequelize | SQLite3 | Socket.IO | WebRTC\nRSA-OAEP 2048-bit | AES-GCM 256-bit | JWT | bcrypt')
r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
doc.add_page_break()

# ════════════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
for item in [
    '1. Project Overview',
    '2. Technology Stack',
    '3. Architecture Diagram',
    '4. Project File Structure',
    '5. Module Descriptions',
    '   5.1 Entry Point — server.js',
    '   5.2 Config Layer',
    '   5.3 Models Layer (Sequelize ORM)',
    '   5.4 Controllers Layer',
    '   5.5 Routes Layer',
    '   5.6 Middleware Layer',
    '   5.7 Sockets Layer',
    '   5.8 Frontend Modules',
    '6. Database Schema (ER Diagram)',
    '7. Encryption Flow',
    '8. Data Flow Diagrams',
    '   8.1 DFD Level 0 — Context Diagram',
    '   8.2 DFD Level 1 — Major Subsystems',
    '   8.3 DFD Level 2 — Authentication System',
    '   8.4 DFD Level 2 — Messaging System',
    '   8.5 DFD Level 2 — Credit Management',
    '9. API Endpoints Reference',
    '10. Socket.IO Events Reference',
    '11. Security Model Summary',
    '12. Key Data Flows — Plain English',
]:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
doc.add_page_break()

# ════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ════════════════════════════════════════════════
doc.add_heading('1. Project Overview', level=1)
doc.add_paragraph(
    'SecureConnect is an end-to-end encrypted (E2EE) real-time messaging and voice-calling '
    'web application built with a modern MVC (Model–View–Controller) architecture. All cryptographic '
    'operations happen on the client side — the server never sees plaintext messages or private keys.'
)
doc.add_paragraph(
    'The application was recently refactored from a monolithic server.js into a modular MVC '
    'structure with Sequelize ORM, separated controllers, routes, middleware, socket event handlers, '
    'and a frontend module system. Key improvements include database transactions for data integrity, '
    'optional Redis adapter for horizontal scaling, proper CSP configuration, automated tests, and '
    'new informational pages (About, Contact, 404).'
)

add_table(['Attribute', 'Detail'], [
    ['Type', 'Full-Stack Web Application'],
    ['Architecture', 'MVC (Model–View–Controller)'],
    ['Backend', 'Node.js + Express.js'],
    ['ORM', 'Sequelize (SQLite3 dialect)'],
    ['Frontend', 'Vanilla HTML5 / CSS3 / JavaScript (ES6 Modules)'],
    ['Database', 'SQLite3 (file-based relational DB)'],
    ['Real-Time', 'Socket.IO (with optional Redis Adapter for scaling)'],
    ['Voice Calls', 'WebRTC (Peer-to-Peer)'],
    ['Encryption', 'RSA-OAEP 2048-bit + AES-GCM 256-bit'],
    ['Authentication', 'JWT + bcrypt + httpOnly Cookies'],
    ['Testing', 'Jest + Supertest + Playwright'],
    ['Deployment', 'AWS EC2 via GitHub Actions CI/CD + PM2'],
])

# ════════════════════════════════════════════════
# 2. TECHNOLOGY STACK
# ════════════════════════════════════════════════
doc.add_heading('2. Technology Stack', level=1)
doc.add_heading('Backend', level=2)
add_table(['Technology', 'Purpose'], [
    ['Node.js', 'Server-side JavaScript runtime'],
    ['Express.js', 'HTTP server framework, REST API routing'],
    ['Sequelize', 'ORM for database abstraction, model definitions, associations, transactions'],
    ['SQLite3', 'Lightweight file-based relational database'],
    ['Socket.IO', 'WebSocket server for real-time bidirectional events'],
    ['@socket.io/redis-adapter + Redis', 'Optional multi-instance Socket.IO sync for horizontal scaling'],
    ['bcryptjs', 'Password hashing with salt (10 rounds)'],
    ['jsonwebtoken', 'JWT creation and verification'],
    ['helmet', 'HTTP security headers with CSP configured'],
    ['express-rate-limit', 'Brute-force protection on auth endpoints'],
    ['multer', 'Multipart file upload handling (avatars)'],
    ['dotenv', 'Environment variable loading from .env file'],
    ['cookie-parser + express-session', 'Cookie/session management'],
])

doc.add_heading('Frontend', level=2)
add_table(['Technology', 'Purpose'], [
    ['HTML5', 'Page structure (7 pages: index, login, chat, admin, about, contact, 404)'],
    ['CSS3 (Vanilla)', 'Styling with 6 stylesheets (landing, style, chat, admin, about, contact)'],
    ['Vanilla JavaScript', 'Client-side logic with modular ES6 architecture'],
    ['Web Crypto API', 'RSA key generation, AES encryption/decryption (E2EE)'],
    ['WebRTC', 'Peer-to-peer voice calling (private & mesh group calls)'],
    ['Socket.IO Client', 'Real-time WebSocket communication'],
])

doc.add_heading('DevOps & Testing', level=2)
add_table(['Technology', 'Purpose'], [
    ['Jest', 'Unit/integration test framework'],
    ['Supertest', 'HTTP assertion library for API testing'],
    ['Playwright', 'End-to-end browser testing'],
    ['cross-env', 'Cross-platform environment variable setting for tests'],
    ['GitHub Actions', 'CI/CD pipeline for automated deployment'],
    ['PM2', 'Node.js process manager (auto-restart, monitoring)'],
    ['AWS EC2', 'Cloud hosting for production deployment'],
])

doc.add_page_break()

# ════════════════════════════════════════════════
# 3. ARCHITECTURE DIAGRAM
# ════════════════════════════════════════════════
doc.add_heading('3. Architecture Diagram', level=1)
doc.add_paragraph('The following diagram shows the layered MVC architecture of SecureConnect:')

mono("""
┌─────────────────────────────────────────────────────────────────────┐
│                         CLIENT (Browser)                            │
│                                                                     │
│  ┌──────────┐  ┌───────────┐  ┌────────────────┐  ┌─────────────┐  │
│  │ 7 HTML   │  │ 6 CSS     │  │ JS Core        │  │ JS Modules  │  │
│  │ Pages    │  │ Sheets    │  │                │  │ (new)       │  │
│  │          │  │           │  │ auth.js        │  │             │  │
│  │ index    │  │ landing   │  │ chat.js        │  │ api.js      │  │
│  │ login    │  │ style     │  │ crypto.js      │  │ socket.js   │  │
│  │ chat     │  │ chat      │  │ admin.js       │  │ ui.js       │  │
│  │ admin    │  │ admin     │  │ contact.js     │  │ webrtc.js   │  │
│  │ about    │  │ about     │  │ index.js       │  │ adminApi.js │  │
│  │ contact  │  │ contact   │  │                │  │             │  │
│  │ 404      │  │           │  │ Web Crypto API │  │ WebRTC      │  │
│  └──────────┘  └───────────┘  └────────────────┘  └─────────────┘  │
│                        │ HTTP REST  │  Socket.IO  │  WebRTC (P2P)  │
└────────────────────────┼────────────┼─────────────┼────────────────┘
                         │            │             │
                         ▼            ▼             │
┌─────────────────────────────────────────────────────────────────────┐
│                     server.js (Entry Point - 119 lines)             │
│    Express setup ─ Helmet/CSP ─ Sessions ─ Route mounting           │
└──────────────────────────────┬──────────────────────────────────────┘
                               │
          ┌────────────────────┼────────────────────┐
          │                    │                    │
          ▼                    ▼                    ▼
┌─────────────────┐  ┌─────────────────┐  ┌─────────────────────────┐
│  MIDDLEWARE      │  │  ROUTES          │  │  SOCKETS                │
│                  │  │                  │  │                         │
│  auth.js         │  │  authRoutes      │  │  socketManager.js       │
│  (JWT + ban)     │  │  userRoutes      │  │  (Redis adapter opt.)   │
│                  │  │  messageRoutes   │  │                         │
│  rateLimit.js    │  │  creditRoutes    │  │  events/                │
│  asyncHandler.js │  │  adminRoutes     │  │    chatEvents.js        │
│  errorHandler.js │  │  groupRoutes     │  │    groupEvents.js       │
│                  │  │                  │  │    webrtcEvents.js      │
└─────────────────┘  └──────┬───────────┘  │                         │
                            │              │  utils/                  │
                            ▼              │    creditHelper.js       │
                   ┌─────────────────┐     └─────────────────────────┘
                   │  CONTROLLERS     │
                   │                  │
                   │  authController  │
                   │  userController  │
                   │  messageController│
                   │  creditController│
                   │  adminController │
                   │  groupController │
                   └──────┬───────────┘
                          │
                          ▼
                   ┌─────────────────┐     ┌───────────────────────┐
                   │  MODELS          │     │  CONFIG               │
                   │  (Sequelize ORM) │     │                       │
                   │                  │     │  env.js (dotenv)      │
                   │  User            │     │  db.js (Sequelize     │
                   │  Message         │     │         init + seed)  │
                   │  Group           │     │  upload.js (Multer)   │
                   │  GroupMember     │     │                       │
                   │  GroupMessage    │     └───────────────────────┘
                   │  CreditTransaction│
                   └──────┬───────────┘
                          │
                          ▼
                   ┌─────────────────┐
                   │  SQLite3 DB      │
                   │  secureconnect.db│
                   └─────────────────┘
""", 7)

doc.add_page_break()

# ════════════════════════════════════════════════
# 4. FILE STRUCTURE
# ════════════════════════════════════════════════
doc.add_heading('4. Project File Structure', level=1)
mono("""SecureConnect/
├── server.js                  ← Thin entry point (119 lines)
├── .env                       ← Environment variables
├── package.json               ← Dependencies & scripts
├── secureconnect.db           ← SQLite database (auto-created)
├── 4-Phases.md                ← Refactoring recommendations
│
├── src/                       ← Backend MVC source code
│   ├── config/
│   │   ├── env.js             ← dotenv loader
│   │   ├── db.js              ← Sequelize init, admin seeding
│   │   └── upload.js          ← Multer config (2MB, images)
│   ├── models/
│   │   ├── index.js           ← Sequelize instance + associations
│   │   ├── User.js            ├── Message.js
│   │   ├── Group.js           ├── GroupMember.js
│   │   ├── GroupMessage.js    └── CreditTransaction.js
│   ├── controllers/
│   │   ├── authController.js  ← Register, Login, Logout, Key Sync
│   │   ├── userController.js  ← Profile, Avatar, Search
│   │   ├── messageController.js ← Chats, History
│   │   ├── creditController.js  ← Request, Transactions
│   │   ├── adminController.js   ← User mgmt, Approvals
│   │   └── groupController.js   ← Full group CRUD
│   ├── routes/
│   │   ├── authRoutes.js      ├── userRoutes.js
│   │   ├── messageRoutes.js   ├── creditRoutes.js
│   │   ├── adminRoutes.js     └── groupRoutes.js
│   ├── middleware/
│   │   ├── auth.js            ← JWT + ban check
│   │   ├── rateLimit.js       ← 10 req / 15 min
│   │   ├── asyncHandler.js    └── errorHandler.js
│   └── sockets/
│       ├── socketManager.js   ← Redis adapter, auth, cleanup
│       ├── events/
│       │   ├── chatEvents.js  ├── groupEvents.js
│       │   └── webrtcEvents.js
│       └── utils/creditHelper.js
│
├── public/                    ← Frontend static files
│   ├── index.html  login.html  chat.html  admin.html
│   ├── about.html  contact.html  404.html
│   ├── css/ (6 files)  images/ (logos, favicons)
│   ├── js/
│   │   ├── auth.js  chat.js  crypto.js  admin.js
│   │   ├── contact.js  index.js
│   │   └── modules/ (api.js, socket.js, ui.js, webrtc.js, adminApi.js)
│   └── uploads/
│
├── tests/ (auth.test.js, setup.js)
└── .github/workflows/deploy.yml
""", 8)

doc.add_page_break()

# ════════════════════════════════════════════════
# 5. MODULE DESCRIPTIONS
# ════════════════════════════════════════════════
doc.add_heading('5. Module Descriptions', level=1)

doc.add_heading('5.1 Entry Point — server.js', level=2)
doc.add_paragraph(
    'The entry point is now a thin 119-line file that sets up Express with Helmet (CSP configured), '
    'session management, static file serving, mounts all route modules, initializes the database via '
    'Sequelize, and starts Socket.IO. All business logic has been moved to the src/ directory.'
)

doc.add_heading('5.2 Config Layer', level=2)
add_table(['File', 'Purpose'], [
    ['env.js', 'Loads .env via dotenv. Exports PORT, JWT_SECRET, ADMIN_PASSWORD, NODE_ENV. Exits if secrets missing.'],
    ['db.js', 'Initializes Sequelize, syncs models, creates default admin user (bcrypt hashed), resets all login states on boot.'],
    ['upload.js', 'Configures Multer for avatar uploads: 2MB limit, image-only MIME filter, unique filenames.'],
])

doc.add_heading('5.3 Models Layer (Sequelize ORM)', level=2)
add_table(['Model', 'Table', 'Key Fields'], [
    ['User', 'users', 'id, username(unique), password(bcrypt), email(unique), credits(500), public_key, encrypted_private_key, is_admin, is_banned, is_logged_in, avatar'],
    ['Message', 'messages', 'id, from_user(FK), to_user(FK), encrypted_content, is_read, created_at'],
    ['Group', 'groups', 'id, name, description, created_by(FK), created_at'],
    ['GroupMember', 'group_members', 'group_id(FK), user_id(FK), role(admin/member), encrypted_group_key, last_read_at, joined_at'],
    ['GroupMessage', 'group_messages', 'id, group_id(FK), from_user(FK), encrypted_content, created_at'],
    ['CreditTransaction', 'credit_transactions', 'id, user_id(FK), amount, transaction_ref, status(pending/approved/rejected), created_at'],
])
doc.add_paragraph('Associations are defined in models/index.js using Sequelize\'s hasMany, belongsTo, belongsToMany relationships with proper foreign keys.')

doc.add_heading('5.4 Controllers Layer', level=2)
add_table(['Controller', 'Functions', 'Key Features'], [
    ['authController', 'register, login, logout, syncKey', 'Input validation (regex), bcrypt hashing, JWT issuance, httpOnly cookies, Sequelize error handling'],
    ['userController', 'getMe, uploadAvatar, getUsers, searchUsers', 'Attribute-specific queries, group member exclusion in search'],
    ['messageController', 'getRecentChats, getChatHistory', 'Raw SQL via sequelize.query for complex joins, auto mark-as-read'],
    ['creditController', 'requestCredits, getTransactions', 'Simple Sequelize CRUD for credit requests'],
    ['adminController', 'getPendingTransactions, approveTransaction, rejectTransaction, getUsers, banUser, deleteUser', 'Uses sequelize.transaction() for atomic credit approval'],
    ['groupController', 'createGroup, getGroups, getGroupMembers, addGroupMember, removeGroupMember, getGroupMessages, getGroupStatus, markGroupRead', 'Sequelize transactions for create/add/remove, auto-delete empty groups'],
])

doc.add_heading('5.5 Routes Layer', level=2)
add_table(['Router', 'Base Path', 'Middleware'], [
    ['authRoutes', '/api', 'authLimiter on register/login'],
    ['userRoutes', '/api', 'authenticateToken + requireNoAdmin'],
    ['messageRoutes', '/api', 'authenticateToken + requireNoAdmin'],
    ['creditRoutes', '/api/credits', 'authenticateToken + requireNoAdmin'],
    ['adminRoutes', '/api/admin', 'authenticateToken + requireAdmin (applied to all)'],
    ['groupRoutes', '/api/groups', 'authenticateToken (applied to all)'],
])

doc.add_heading('5.6 Middleware Layer', level=2)
add_table(['Middleware', 'Purpose'], [
    ['auth.js', 'Extracts JWT from cookie or Authorization header, verifies signature, checks ban status via Sequelize DB lookup, provides requireAdmin and requireNoAdmin guards'],
    ['rateLimit.js', 'Limits auth endpoints to 10 requests per 15 minutes per IP'],
    ['asyncHandler.js', 'Wraps async route handlers to forward errors to Express error middleware'],
    ['errorHandler.js', 'Global error handler: returns HTML for browser requests, JSON for API requests, hides stack trace in production'],
])

doc.add_heading('5.7 Sockets Layer', level=2)
add_table(['Module', 'Purpose'], [
    ['socketManager.js', 'Initializes Socket.IO with optional Redis adapter for multi-instance scaling. Handles authentication, connected user tracking, disconnect cleanup, online status broadcasts.'],
    ['chatEvents.js', 'Handles send-message event: credit deduction via creditHelper, Message.create() via Sequelize, forwarding to recipient socket'],
    ['groupEvents.js', 'Handles join-group and send-group-message: membership verification, credit deduction, GroupMessage.create(), broadcast to group room'],
    ['webrtcEvents.js', 'All WebRTC signaling: call-request, call-response, ice-candidate, call-ended, group-call-request/join/leave, group-call-offer/answer/ice, call-connected credit deduction'],
    ['creditHelper.js', 'Atomic credit deduction using User.decrement() with Sequelize WHERE guard (credits > 0)'],
])

doc.add_heading('5.8 Frontend Modules', level=2)
add_table(['Module', 'Purpose'], [
    ['auth.js (286 lines)', 'Login/register forms, RSA key generation on registration, JWT storage, session validation, logout'],
    ['chat.js (37KB)', 'Chat UI, messaging, groups, credit modal — now uses imported modules'],
    ['crypto.js (447 lines)', 'SecureEncryption class: RSA-OAEP keygen, AES-GCM encrypt/decrypt, PBKDF2 key derivation, group key management'],
    ['admin.js', 'Admin login, user management dashboard, credit transaction review'],
    ['modules/api.js', 'REST API helper functions (fetch wrappers with auth headers)'],
    ['modules/socket.js', 'Socket.IO connection module (extracted from chat.js)'],
    ['modules/ui.js', 'UI rendering utilities (extracted DOM manipulation)'],
    ['modules/webrtc.js', 'WebRTC call module (private + mesh group, separated from chat.js)'],
    ['modules/adminApi.js', 'Admin API helper functions'],
    ['contact.js', 'Contact form handler'],
    ['index.js', 'Landing page interactions'],
])

doc.add_page_break()

# ════════════════════════════════════════════════
# 6. DATABASE SCHEMA
# ════════════════════════════════════════════════
doc.add_heading('6. Database Schema', level=1)

doc.add_heading('Users Table', level=3)
add_table(['Column', 'Type', 'Description'], [
    ['id', 'INTEGER PK', 'Auto-increment primary key'],
    ['username', 'STRING UNIQUE', 'Lowercase alphanumeric (5-32 chars)'],
    ['password', 'STRING', 'bcrypt-hashed password'],
    ['email', 'STRING UNIQUE', 'User email address'],
    ['credits', 'INTEGER', 'Credit balance (default: 500)'],
    ['public_key', 'TEXT', 'RSA-2048 public key (Base64)'],
    ['encrypted_private_key', 'TEXT', 'Password-encrypted RSA private key'],
    ['is_logged_in', 'INTEGER', 'Online status flag (0/1)'],
    ['is_admin', 'INTEGER', 'Admin flag (0/1)'],
    ['is_banned', 'INTEGER', 'Ban status (0/1)'],
    ['avatar', 'TEXT', 'Path to avatar image'],
    ['created_at', 'DATE', 'Account creation timestamp'],
])

doc.add_heading('Messages Table', level=3)
add_table(['Column', 'Type', 'Description'], [
    ['id', 'INTEGER PK', 'Auto-increment'], ['from_user', 'INTEGER FK', 'Sender user ID'],
    ['to_user', 'INTEGER FK', 'Recipient user ID'], ['encrypted_content', 'TEXT', 'E2EE encrypted blob'],
    ['is_read', 'INTEGER', 'Read status (0/1)'], ['created_at', 'DATE', 'Timestamp'],
])

doc.add_heading('Groups Table', level=3)
add_table(['Column', 'Type', 'Description'], [
    ['id', 'INTEGER PK', 'Auto-increment'], ['name', 'STRING', 'Group name'],
    ['description', 'TEXT', 'Description'], ['created_by', 'INTEGER FK', 'Creator user ID'],
    ['created_at', 'DATE', 'Timestamp'],
])

doc.add_heading('Group Members Table', level=3)
add_table(['Column', 'Type', 'Description'], [
    ['group_id', 'INTEGER FK', 'Group reference'], ['user_id', 'INTEGER FK', 'User reference'],
    ['role', 'STRING', '"admin" or "member"'], ['encrypted_group_key', 'TEXT', 'RSA-encrypted group AES key'],
    ['last_read_at', 'DATE', 'Last read timestamp'], ['joined_at', 'DATE', 'Join timestamp'],
])

doc.add_heading('Group Messages Table', level=3)
add_table(['Column', 'Type', 'Description'], [
    ['id', 'INTEGER PK', 'Auto-increment'], ['group_id', 'INTEGER FK', 'Group reference'],
    ['from_user', 'INTEGER FK', 'Sender user ID'], ['encrypted_content', 'TEXT', 'Encrypted blob'],
    ['created_at', 'DATE', 'Timestamp'],
])

doc.add_heading('Credit Transactions Table', level=3)
add_table(['Column', 'Type', 'Description'], [
    ['id', 'INTEGER PK', 'Auto-increment'], ['user_id', 'INTEGER FK', 'Requesting user'],
    ['amount', 'INTEGER', 'Requested amount'], ['transaction_ref', 'STRING', 'Payment reference'],
    ['status', 'STRING', 'pending/approved/rejected'], ['created_at', 'DATE', 'Timestamp'],
])

doc.add_heading('Entity Relationships', level=3)
doc.add_paragraph('• Users (1) ── sends/receives ──→ (Many) Messages')
doc.add_paragraph('• Users (1) ── creates ──→ (Many) Groups')
doc.add_paragraph('• Users (Many) ←─ belongs via GroupMembers ─→ (Many) Groups')
doc.add_paragraph('• Groups (1) ── contains ──→ (Many) Group Messages')
doc.add_paragraph('• Users (1) ── requests ──→ (Many) Credit Transactions')

doc.add_page_break()

# ════════════════════════════════════════════════
# 7. ENCRYPTION FLOW
# ════════════════════════════════════════════════
doc.add_heading('7. Encryption Flow (How E2EE Works)', level=1)

doc.add_heading('Registration (Key Generation)', level=2)
doc.add_paragraph('1. User fills registration form (username, email, password).')
doc.add_paragraph('2. Browser generates RSA-2048 key pair using Web Crypto API.')
doc.add_paragraph('3. Public key exported to Base64 string format.')
doc.add_paragraph('4. Private key encrypted using AES-GCM with a key derived from password (PBKDF2 with random salt).')
doc.add_paragraph('5. Both keys sent to server. Server stores them via Sequelize User.create().')
doc.add_paragraph('6. Private key is NEVER stored in plaintext on the server.')

doc.add_heading('Sending a Private Message', level=2)
doc.add_paragraph('1. Sender types message. Browser generates random AES-256 key.')
doc.add_paragraph('2. Message encrypted with AES-GCM using this key.')
doc.add_paragraph('3. AES key encrypted with recipient\'s RSA public key AND sender\'s RSA public key (for history).')
doc.add_paragraph('4. Encrypted package sent via Socket.IO → chatEvents.js deducts credit → Message.create() stores blob.')
doc.add_paragraph('5. Server forwards to recipient socket. Cannot decrypt the content.')

doc.add_heading('Receiving / Decrypting', level=2)
doc.add_paragraph('1. Recipient receives encrypted package via Socket.IO.')
doc.add_paragraph('2. Decrypts AES key with own RSA private key.')
doc.add_paragraph('3. Decrypts message with AES key. Displays plaintext in chat UI.')

doc.add_heading('Group Encryption', level=2)
doc.add_paragraph('1. Group admin generates communal AES-256 group key.')
doc.add_paragraph('2. Key individually encrypted with each member\'s RSA public key.')
doc.add_paragraph('3. Stored per-member in group_members table via Sequelize.')
doc.add_paragraph('4. All group messages encrypted symmetrically with shared group AES key.')

mono("""
  SENDER (Browser)                SERVER (Node.js)              RECEIVER (Browser)
       │                                │                              │
  1. Generate AES-256 key               │                              │
  2. Encrypt msg (AES-GCM)              │                              │
  3. Encrypt AES key with               │                              │
     Receiver's RSA pub key             │                              │
       │                                │                              │
       ├── Socket.IO "send-message" ──→ │                              │
       │                                ├── creditHelper.deductCredit()│
       │                                ├── Message.create() (Sequelize)
       │                                ├── Forward via Socket.IO ────→│
       │                                │                              │
       │                                │                 4. Decrypt AES key
       │                                │                    with RSA priv key
       │                                │                 5. Decrypt message
       │                                │                    with AES key
       │                                │                 6. Display plaintext
""", 8)

doc.add_page_break()

# ════════════════════════════════════════════════
# 8. DATA FLOW DIAGRAMS
# ════════════════════════════════════════════════
doc.add_heading('8. Data Flow Diagrams', level=1)

# DFD Level 0
doc.add_heading('8.1 DFD Level 0 — Context Diagram', level=2)
doc.add_paragraph('Shows SecureConnect as a single process with external entities and data stores.')
mono("""
 ┌──────────┐                                           ┌──────────┐
 │          │  Register / Login / Send Messages /        │          │
 │   USER   │  Make Calls / Request Credits              │          │
 │          │ ─────────────────────────────────────────→ │          │
 │(External │                                            │ SECURE   │
 │ Entity)  │ ←──────────────────────────────────────── │ CONNECT  │
 │          │  Encrypted Messages / Call Signals /       │ SYSTEM   │
 └──────────┘  Credit Balance / Notifications            │          │
                                                         │(Process) │
 ┌──────────┐  Login / Manage Users /                    │          │
 │  ADMIN   │  Approve-Reject Credits                    │          │
 │(External │ ─────────────────────────────────────────→ │          │
 │ Entity)  │ ←──────────────────────────────────────── │          │
 └──────────┘  User List / Transaction List              └────┬─────┘
                                                              │
                    ┌────────────────────────────────┐        │
                    │  SQLite Database (Sequelize)    │◄──────►│
                    │  secureconnect.db               │
                    │  6 tables: users, messages,     │
                    │  groups, group_members,          │
                    │  group_messages, credit_transactions │
                    └────────────────────────────────┘
""", 7.5)

doc.add_page_break()

# DFD Level 1
doc.add_heading('8.2 DFD Level 1 — Major Subsystems', level=2)
doc.add_paragraph('Decomposes the system into 5 core processes mapped to controller files.')
mono("""
 ┌──────────┐                                              ┌──────────┐
 │   USER   │                                              │  ADMIN   │
 └────┬─────┘                                              └────┬─────┘
      │                                                         │
      │ Credentials + Keys       ┌──────────────────────────────┘
      │                          │ Admin Credentials
      ▼                          ▼
 ┌───────────────────────────────────────┐
 │  1.0 AUTHENTICATION SYSTEM            │  ←─── authController.js
 │  Register, Login, Logout, Key Sync    │       authRoutes.js
 └──┬──────────────────────────────┬─────┘
    │ JWT Token                    │ Sequelize CRUD    ┌────────────┐
    │ + Profile                    └──────────────────→│ D1: Users  │
    │                                                  └──────┬─────┘
    │ ─────────────────────────────────────────────────────────┤
    ▼                                                         │
 ┌──────────────────────┐   ┌──────────────────────┐          │
 │ 2.0 MESSAGING SYSTEM │   │ 3.0 VOICE CALLING    │          │
 │                      │   │ SYSTEM                │          │
 │ messageController +  │   │                       │          │
 │ chatEvents +         │◄──┤ webrtcEvents.js       │◄─────────┤
 │ groupEvents          │   │                       │ deductCredit
 └──────────┬───────────┘   └───────────────────────┘          │
            │ Msg CRUD                                         │
            ▼                                                  │
  ┌──────────────┐  ┌──────────────────┐                       │
  │ D2: Messages │  │ D3: Groups &     │                       │
  │ +GroupMessages│  │     Members      │                       │
  └──────────────┘  └──────────────────┘                       │
                                                               │
 ┌──────────────────────┐              ┌──────────────────────┐│
 │ 4.0 CREDIT MGMT      │              │ 5.0 ADMIN MGMT      ││
 │ creditController.js   │              │ adminController.js   ││
 └──────────┬────────────┘              └───────────┬──────────┘│
            │                                       │           │
            ▼                                       ▼           │
  ┌───────────────────┐               ┌─────────────────┐      │
  │ D4: Credit        │◄─────────────►│ D1: Users       │◄─────┘
  │     Transactions  │               └─────────────────┘
  └───────────────────┘
""", 7)

add_table(['Process', 'MVC Files', 'Description'], [
    ['1.0 Authentication', 'authController.js + authRoutes.js', 'Register with RSA keys, Login with JWT + httpOnly cookie, Logout, Key Sync endpoint'],
    ['2.0 Messaging', 'messageController + chatEvents + groupEvents', 'Real-time E2EE private & group messaging via Socket.IO + Sequelize models'],
    ['3.0 Voice Calling', 'webrtcEvents.js', 'WebRTC signaling, private P2P calls, mesh group calls, credit deduction on connect'],
    ['4.0 Credit Mgmt', 'creditController.js + creditRoutes.js', 'Credit request submission, transaction history viewing'],
    ['5.0 Admin Mgmt', 'adminController.js + adminRoutes.js', 'User ban/unban/delete, transaction approve/reject with Sequelize DB transactions'],
])

doc.add_page_break()

# DFD Level 2 — Auth
doc.add_heading('8.3 DFD Level 2 — Process 1.0: Authentication System', level=2)
mono("""
 ┌──────────┐
 │   USER   │
 └────┬─────┘
      │ Username, Email, Password, Public Key, Encrypted Private Key
      ▼
 ┌───────────────────────────────────────────────────────┐
 │  1.1 REGISTER USER (authController.register)          │
 │                                                       │
 │  • Validate input (username regex, email, password)   │
 │  • bcrypt.hash(password, 10)                          │
 │  • User.create() via Sequelize                        │
 │  • jwt.sign() → httpOnly cookie + response body       │
 │  • Initial credits = 500                              │
 └─────────────────────────┬─────────────────────────────┘
                           │ Sequelize User.create()
                           ▼
                    ┌────────────┐
                    │ D1: Users  │
                    │ (Sequelize)│
                    └──────┬─────┘
    ┌──────────────────────┼────────────────────────┐
    │                      │                        │
    ▼                      ▼                        ▼
┌────────────────┐ ┌────────────────┐  ┌───────────────────┐
│1.2 LOGIN USER  │ │1.3 VALIDATE    │  │1.4 LOGOUT USER    │
│                │ │SESSION         │  │                   │
│authController  │ │                │  │authController     │
│.login          │ │middleware/     │  │.logout            │
│                │ │auth.js         │  │                   │
│• User.findOne()│ │• jwt.verify()  │  │• User.update(     │
│• bcrypt.compare│ │• User.findByPk │  │  is_logged_in=0)  │
│• Check ban/    │ │  (ban check)   │  │• Clear cookie     │
│  login status  │ │                │  │                   │
│• jwt.sign()    │ └────────────────┘  └───────────────────┘
│• Set cookie    │
└────────────────┘
                        ┌────────────────┐
                        │1.5 SYNC KEYS   │
                        │                │
                        │authController  │
                        │.syncKey        │
                        │                │
                        │• User.update(  │
                        │  encrypted_    │
                        │  private_key)  │
                        └────────────────┘
""", 7.5)

doc.add_page_break()

# DFD Level 2 — Messaging
doc.add_heading('8.4 DFD Level 2 — Process 2.0: Messaging System', level=2)
mono("""
 ┌──────────┐
 │   USER   │
 └────┬─────┘
      │
      ├── Encrypted content + Recipient ID (Socket.IO)
      │                      │
      │                      ▼
      │              ┌────────────────────────────┐
      │              │ 2.1 SEND PRIVATE MESSAGE    │
      │              │ (chatEvents.js)              │     ┌──────────┐
      │              │                              │────→│D1: Users │
      │              │ • creditHelper.deductCredit()│     └──────────┘
      │              │ • Message.create() (Sequelize)│───→┌──────────┐
      │              │ • Forward to recipient socket │    │D2: Msgs  │
      │              └──────────────┬────────────────┘    └──────────┘
      │                             │ Socket.IO forward
      │                             ▼
      │              ┌────────────────────────────┐
      │              │ 2.2 RECEIVE MESSAGE         │
      │              │ Deliver encrypted packet    │─────→ USER
      │              └────────────────────────────┘
      │
      ├── Encrypted content + Group ID (Socket.IO)
      │                      ▼
      │              ┌────────────────────────────┐     ┌──────────────┐
      │              │ 2.3 SEND GROUP MESSAGE      │────→│D3: Groups &  │
      │              │ (groupEvents.js)             │    │   Members    │
      │              │ • GroupMember.findOne()       │    └──────────────┘
      │              │ • creditHelper.deductCredit() │
      │              │ • GroupMessage.create()       │
      │              │ • Broadcast to group room     │
      │              └────────────────────────────┘
      │
      ├── Request chat history (HTTP)
      │                      ▼
      │              ┌────────────────────────────┐
      │              │ 2.4 LOAD CHAT HISTORY       │
      │              │ (messageController)          │◄──→ D2: Messages
      │              │ • Message.findAll() + joins   │
      │              │ • Message.update(is_read=1)   │
      │              └────────────────────────────┘
      │
      └── Create/Add/Remove group (HTTP)
                             ▼
                     ┌────────────────────────────┐
                     │ 2.5 MANAGE GROUPS           │
                     │ (groupController)            │◄──→ D3: Groups
                     │ • sequelize.transaction()    │
                     │ • Group/GroupMember CRUD      │
                     │ • Auto-delete empty groups    │
                     └────────────────────────────┘
""", 7)

doc.add_page_break()

# DFD Level 2 — Credit
doc.add_heading('8.5 DFD Level 2 — Process 4.0: Credit Management', level=2)
mono("""
 ┌──────────┐                                         ┌──────────┐
 │   USER   │                                         │  ADMIN   │
 └────┬─────┘                                         └────┬─────┘
      │ Amount +                           Approve         │
      │ Transaction Ref                   Transaction ID   │
      ▼                                                    ▼
 ┌─────────────────────┐              ┌──────────────────────────┐
 │4.1 REQUEST CREDITS  │              │4.3 APPROVE TRANSACTION   │
 │                     │              │    (adminController)     │
 │creditController     │              │                          │
 │.requestCredits      │              │• sequelize.transaction() │
 │                     │              │• CreditTransaction.      │
 │• Validate amount    │              │  findByPk()              │
 │• CreditTransaction  │              │• Update status="approved"│
 │  .create(pending)   │              │• user.credits += amount  │
 └──────────┬──────────┘              └───────────┬──────────────┘
            │                                     │
            ▼                                     ▼
     ┌─────────────────────────────────────────────────┐
     │         D4: Credit Transactions (Sequelize)     │
     └──────────────┬──────────────────────┬───────────┘
                    │                      │
                    ▼                      ▼
 ┌─────────────────────┐              ┌──────────────────────────┐
 │4.2 VIEW TRANSACTIONS│              │4.4 REJECT TRANSACTION    │
 │                     │              │                          │
 │creditController     │              │• CreditTransaction       │
 │.getTransactions     │              │  .update(rejected)       │
 │                     │              │• No credits added        │
 │• CreditTransaction  │              │                          │
 │  .findAll()         │              └──────────────────────────┘
 └─────────────────────┘
""", 7.5)

doc.add_page_break()

# ════════════════════════════════════════════════
# 9. API ENDPOINTS
# ════════════════════════════════════════════════
doc.add_heading('9. API Endpoints Reference', level=1)

doc.add_heading('Authentication & User Routes', level=2)
add_table(['Method', 'Endpoint', 'Purpose', 'Auth'], [
    ['POST', '/api/register', 'User registration + key exchange', 'Rate-limited'],
    ['POST', '/api/login', 'User login, JWT issuance', 'Rate-limited'],
    ['POST', '/api/auth/logout', 'Logout user, clear cookie', 'JWT'],
    ['POST', '/api/user/sync-key', 'Sync encrypted private key', 'JWT (User)'],
    ['GET', '/api/user/me', 'Get current user profile', 'JWT'],
    ['POST', '/api/user/avatar', 'Upload avatar image', 'JWT (User)'],
    ['GET', '/api/users', 'List all non-admin users', 'JWT (User)'],
    ['GET', '/api/users/search', 'Search users by name', 'JWT (User)'],
])

doc.add_heading('Messaging & Chat Routes', level=2)
add_table(['Method', 'Endpoint', 'Purpose', 'Auth'], [
    ['GET', '/api/chats', 'Recent chats with unread counts', 'JWT (User)'],
    ['GET', '/api/messages/:userId', 'Chat history with user', 'JWT (User)'],
])

doc.add_heading('Group Routes', level=2)
add_table(['Method', 'Endpoint', 'Purpose', 'Auth'], [
    ['POST', '/api/groups', 'Create group', 'JWT'],
    ['GET', '/api/groups', 'List user\'s groups', 'JWT'],
    ['GET', '/api/groups/:id/members', 'Group members', 'JWT'],
    ['POST', '/api/groups/:id/members', 'Add member', 'JWT'],
    ['DELETE', '/api/groups/:id/members/:userId', 'Remove/Leave', 'JWT'],
    ['GET', '/api/groups/:id/messages', 'Group messages', 'JWT'],
    ['GET', '/api/groups/:id/status', 'Online status', 'JWT'],
    ['POST', '/api/groups/:id/read', 'Mark as read', 'JWT'],
])

doc.add_heading('Credit Routes', level=2)
add_table(['Method', 'Endpoint', 'Purpose', 'Auth'], [
    ['POST', '/api/credits/request', 'Submit credit request', 'JWT (User)'],
    ['GET', '/api/credits/transactions', 'View own transactions', 'JWT (User)'],
])

doc.add_heading('Admin Routes', level=2)
add_table(['Method', 'Endpoint', 'Purpose', 'Auth'], [
    ['GET', '/api/admin/transactions/pending', 'Pending credit requests', 'Admin'],
    ['POST', '/api/admin/transactions/:id/approve', 'Approve request', 'Admin'],
    ['POST', '/api/admin/transactions/:id/reject', 'Reject request', 'Admin'],
    ['GET', '/api/admin/users', 'All users', 'Admin'],
    ['POST', '/api/admin/users/:id/ban', 'Ban/Unban', 'Admin'],
    ['DELETE', '/api/admin/users/:id', 'Delete user', 'Admin'],
])

# ════════════════════════════════════════════════
# 10. SOCKET.IO EVENTS
# ════════════════════════════════════════════════
doc.add_heading('10. Socket.IO Events Reference', level=1)
add_table(['Event', 'Direction', 'Handler', 'Purpose'], [
    ['authenticate', 'Client→Server', 'socketManager', 'Authenticate socket with JWT'],
    ['send-message', 'Client→Server', 'chatEvents', 'Send encrypted private message'],
    ['receive-message', 'Server→Client', 'chatEvents', 'Deliver encrypted message'],
    ['message-sent', 'Server→Client', 'chatEvents', 'Confirm send success'],
    ['users-online', 'Server→All', 'socketManager', 'Broadcast online user list'],
    ['join-group', 'Client→Server', 'groupEvents', 'Join group socket room'],
    ['send-group-message', 'Client→Server', 'groupEvents', 'Send encrypted group message'],
    ['receive-group-message', 'Server→Group', 'groupEvents', 'Deliver group message'],
    ['call-request', 'Client→Server', 'webrtcEvents', 'Initiate private call'],
    ['call-response', 'Client→Server', 'webrtcEvents', 'Answer/reject call'],
    ['ice-candidate', 'Bidirectional', 'webrtcEvents', 'WebRTC ICE negotiation'],
    ['call-ended', 'Bidirectional', 'webrtcEvents', 'End call notification'],
    ['call-connected', 'Client→Server', 'webrtcEvents', 'Credit deduction trigger'],
    ['group-call-request', 'Client→Server', 'webrtcEvents', 'Initiate group call'],
    ['join-group-call', 'Client→Server', 'webrtcEvents', 'Join active group call'],
    ['leave-group-call', 'Client→Server', 'webrtcEvents', 'Leave group call'],
    ['group-call-offer/answer', 'Bidirectional', 'webrtcEvents', 'Group call WebRTC signals'],
    ['group-ice-candidate', 'Bidirectional', 'webrtcEvents', 'Group call ICE signals'],
])

# ════════════════════════════════════════════════
# 11. SECURITY MODEL
# ════════════════════════════════════════════════
doc.add_heading('11. Security Model Summary', level=1)
add_table(['Layer', 'Mechanism', 'Details'], [
    ['Transport', 'HTTPS (production)', 'Encrypted client-server channel'],
    ['Authentication', 'JWT + bcrypt + httpOnly cookies', 'Tokens in cookies + headers; 24h expiry; bcrypt 10 rounds'],
    ['Ban Check', 'Middleware DB check', 'Every request checks ban status via Sequelize User.findByPk()'],
    ['Rate Limiting', 'express-rate-limit', '10 auth attempts / 15 min / IP'],
    ['HTTP Headers', 'Helmet + CSP', 'XSS, clickjacking protection; CSP with script-src, style-src, font-src configured'],
    ['Message Privacy', 'RSA-OAEP + AES-GCM', 'True E2EE — server only stores ciphertext'],
    ['Key Management', 'Web Crypto API', 'Private keys generated and used only in browser'],
    ['Key Sync', 'PBKDF2 + AES-GCM', 'Private key encrypted with password for cross-device sync'],
    ['Data Integrity', 'Sequelize transactions', 'Credit approval, group CRUD use atomic DB transactions'],
    ['Input Validation', 'Regex + Sequelize constraints', 'Username, email, password validated; unique constraints enforced'],
    ['Error Handling', 'Global errorHandler', 'Stack traces hidden in production; HTML/JSON responses'],
    ['Env Security', 'dotenv + required vars', 'Server refuses to start without JWT_SECRET and ADMIN_PASSWORD'],
])

# ════════════════════════════════════════════════
# 12. KEY DATA FLOWS
# ════════════════════════════════════════════════
doc.add_heading('12. Key Data Flows — Plain English', level=1)

doc.add_heading('1. Registration Flow', level=2)
doc.add_paragraph(
    'User fills form → Browser generates RSA-2048 keys (Web Crypto API) → Private key encrypted with password '
    '(PBKDF2 + AES-GCM) → Sent to server → authController.register() validates input, hashes password (bcrypt), '
    'calls User.create() via Sequelize → JWT issued as httpOnly cookie + response body → Redirect to chat.'
)

doc.add_heading('2. Login Flow', level=2)
doc.add_paragraph(
    'Submit credentials → authController.login() calls User.findOne() → bcrypt comparison → Ban/session checks → '
    'User.save(is_logged_in=1) → JWT cookie set → Returns profile + encryption keys → Redirect to chat or admin.'
)

doc.add_heading('3. Message Sending Flow', level=2)
doc.add_paragraph(
    'Type message → Generate random AES-256 key → Encrypt message (AES-GCM) → Encrypt AES key with recipient\'s '
    'RSA public key → Emit send-message via Socket.IO → chatEvents.js calls creditHelper.deductCredit() (atomic '
    'Sequelize decrement) → Message.create() stores encrypted blob → Forward to recipient socket → Recipient '
    'decrypts AES key with RSA private key → Decrypts message.'
)

doc.add_heading('4. Voice Call Flow', level=2)
doc.add_paragraph(
    'Click call → call-request Socket.IO event → webrtcEvents.js forwards to recipient → WebRTC offer/answer/ICE '
    'exchange through server → Direct P2P audio stream → call-connected event triggers creditHelper.deductCredit().'
)

doc.add_heading('5. Credit Request Flow', level=2)
doc.add_paragraph(
    'Enter amount + payment ref → creditController.requestCredits() creates record via CreditTransaction.create(pending) '
    '→ Admin sees in dashboard → adminController.approveTransaction() uses sequelize.transaction() to atomically '
    'update status + add credits to user balance.'
)

# ════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'SecureConnect_Project_Analysis.docx')
doc.save(output_path)
print(f"Document saved to: {output_path}")
