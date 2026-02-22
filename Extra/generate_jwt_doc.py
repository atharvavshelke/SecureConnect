"""Generates a Word document explaining JWT token flow in SecureConnect."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2); s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for lvl in range(1, 4):
    doc.styles[f'Heading {lvl}'].font.name = 'Calibri'
    doc.styles[f'Heading {lvl}'].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

def shade(cell, color):
    tc = cell._element.get_or_add_tcPr()
    tc.append(tc.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'}))

def styled_table(headers, rows, color='1e40af'):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.font.bold = True; r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); r.font.size = Pt(10)
        shade(c, color)
    for rd in rows:
        row = t.add_row()
        for i, txt in enumerate(rd):
            row.cells[i].text = txt
            for p in row.cells[i].paragraphs:
                for r in p.runs: r.font.size = Pt(10)
    return t

def code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    p.paragraph_format.left_indent = Cm(1)
    return p

def bold_run(paragraph, text, color=None):
    r = paragraph.add_run(text)
    r.bold = True
    if color: r.font.color.rgb = color
    return r

def normal_run(paragraph, text):
    return paragraph.add_run(text)


# ═══════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
t = doc.add_heading('SecureConnect', level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t2 = doc.add_heading('JWT Token Analysis', level=1)
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('How JWT Tokens Are Created, Stored, and Verified')
r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x66,0x66,0x66)

doc.add_page_break()


# ═══════════════════════════════════════
# SECTION 1: How JWT Tokens Are Made
# ═══════════════════════════════════════
doc.add_heading('1. How JWT Tokens Are Made', level=1)

doc.add_paragraph(
    'JWT tokens are created in one file: src/controllers/authController.js. '
    'They are generated at two points — on Register and on Login.'
)

# --- Register ---
doc.add_heading('1.1 On Register', level=2)
doc.add_paragraph('File: src/controllers/authController.js (line 48-52)')
code_block(
    'const token = jwt.sign(\n'
    '    { id: newUser.id, username, isAdmin: false },\n'
    '    JWT_SECRET,\n'
    '    { expiresIn: \'24h\' }\n'
    ');'
)
doc.add_paragraph(
    'When a new user registers, a token is signed with their new user ID, username, '
    'and isAdmin set to false (regular users cannot be admins at registration).'
)

# --- Login ---
doc.add_heading('1.2 On Login', level=2)
doc.add_paragraph('File: src/controllers/authController.js (line 95-99)')
code_block(
    'const token = jwt.sign(\n'
    '    { id: user.id, username: user.username, isAdmin: user.is_admin === 1 },\n'
    '    JWT_SECRET,\n'
    '    { expiresIn: \'24h\' }\n'
    ');'
)
doc.add_paragraph(
    'On login, the token includes the actual is_admin status from the database, '
    'so admin users get isAdmin: true in their token.'
)

# --- Payload table ---
doc.add_heading('1.3 What\'s Inside the Token (Payload)', level=2)
styled_table(
    ['Field', 'Type', 'Purpose'],
    [
        ['id', 'Integer', 'User\'s database primary key ID'],
        ['username', 'String', 'User\'s username'],
        ['isAdmin', 'Boolean', 'Whether this is an admin account'],
    ],
    '2d3748'
)

doc.add_paragraph()
p = doc.add_paragraph()
bold_run(p, 'JWT_SECRET: ', RGBColor(0xCC,0x00,0x00))
normal_run(p, 'The signing secret comes from src/config/env.js, which reads it from the .env file. '
              'Without this secret, tokens cannot be created or verified.')

doc.add_paragraph()
p = doc.add_paragraph()
bold_run(p, 'Expiry: ', RGBColor(0xCC,0x00,0x00))
normal_run(p, 'All tokens expire after 24 hours ({ expiresIn: \'24h\' }).')


doc.add_page_break()


# ═══════════════════════════════════════
# SECTION 2: How the Token Gets to the Browser
# ═══════════════════════════════════════
doc.add_heading('2. How the Token Gets to the Browser', level=1)

doc.add_paragraph(
    'After creating the token, the server delivers it to the browser in TWO ways simultaneously:'
)

doc.add_heading('2.1 HTTP-Only Cookie (Server-Set)', level=2)
doc.add_paragraph('File: src/controllers/authController.js')
code_block(
    'res.cookie(\'token\', token, {\n'
    '    httpOnly: true,\n'
    '    maxAge: 24 * 60 * 60 * 1000,   // 24 hours\n'
    '    sameSite: \'strict\',\n'
    '    secure: process.env.NODE_ENV === \'production\'\n'
    '});'
)
doc.add_paragraph(
    'This cookie is httpOnly (JavaScript cannot access it), sameSite: strict (prevents CSRF), '
    'and secure in production (HTTPS only). The browser automatically sends it with every request.'
)

doc.add_heading('2.2 JSON Response Body (Client-Stored)', level=2)
doc.add_paragraph('File: src/controllers/authController.js')
code_block('res.json({ token, userId, username, isAdmin, credits, publicKey, encryptedPrivateKey });')
doc.add_paragraph('File: public/js/auth.js (line 100-102)')
code_block(
    'localStorage.setItem(\'token\', data.token);\n'
    'localStorage.setItem(\'userId\', data.userId);\n'
    'localStorage.setItem(\'username\', data.username);'
)
doc.add_paragraph(
    'The token is also sent in the response JSON so the frontend can store it in '
    'localStorage and include it in API requests and Socket.IO authentication.'
)

doc.add_heading('2.3 Why Two Methods?', level=2)
styled_table(
    ['Method', 'Used For', 'Security'],
    [
        ['HTTP Cookie', 'Automatic auth on page navigation (GET /chat)', 'httpOnly — JS can\'t steal it'],
        ['localStorage', 'API calls (Authorization header) + Socket.IO auth', 'Accessible by JS — needed for fetch() and socket.emit()'],
    ],
    '15803d'
)


doc.add_page_break()


# ═══════════════════════════════════════
# SECTION 3: Where the Token Is Verified (3 Places)
# ═══════════════════════════════════════
doc.add_heading('3. Where the Token Is Verified', level=1)

doc.add_paragraph('The token is verified in THREE distinct places in the codebase:')

# --- Place 1 ---
doc.add_heading('3.1 HTTP Requests — Express Middleware', level=2)
doc.add_paragraph('File: src/middleware/auth.js → authenticateToken()')
code_block(
    'const authenticateToken = async (req, res, next) => {\n'
    '    const token = req.cookies.token ||                    // Check cookie first\n'
    '                  req.headers[\'authorization\']?.split(\' \')[1]; // Then Authorization header\n'
    '\n'
    '    if (!token) return res.status(401).json({ error: \'Access denied\' });\n'
    '\n'
    '    const verified = jwt.verify(token, JWT_SECRET);\n'
    '    req.user = verified;  // { id, username, isAdmin }\n'
    '\n'
    '    // Also checks if user is banned\n'
    '    const user = await User.findByPk(verified.id);\n'
    '    if (!user || user.is_banned) return res.status(403).json({ error: \'Banned\' });\n'
    '    next();\n'
    '};'
)

doc.add_paragraph('This middleware protects these routes:')
styled_table(
    ['Route File', 'Protected Endpoints'],
    [
        ['authRoutes.js', 'POST /api/auth/logout, POST /api/user/sync-key'],
        ['userRoutes.js', 'GET /api/user/me, GET /api/users, POST /api/user/profile-pic'],
        ['messageRoutes.js', 'GET /api/messages/:userId'],
        ['creditRoutes.js', 'GET /api/credits/balance, POST /api/credits/purchase'],
        ['adminRoutes.js', 'All /api/admin/* routes (+ requireAdmin check)'],
        ['groupRoutes.js', 'All /api/groups/* routes'],
        ['server.js', 'GET /chat page (requires authenticateToken + requireNoAdmin)'],
    ],
    '1e40af'
)

# --- Place 2 ---
doc.add_paragraph()
doc.add_heading('3.2 WebSocket Connections — Socket.IO', level=2)
doc.add_paragraph('File: src/sockets/socketManager.js')
code_block(
    'socket.on(\'authenticate\', (token) => {\n'
    '    const verified = jwt.verify(token, JWT_SECRET);\n'
    '    socket.userId = verified.id;\n'
    '    socket.username = verified.username;\n'
    '    socket.isAdmin = verified.isAdmin;\n'
    '    connectedUsers.set(verified.id, socket.id);\n'
    '});'
)
doc.add_paragraph(
    'This is INDEPENDENT of Express — it runs on pure Node.js via Socket.IO. '
    'The token is verified using the same JWT_SECRET but through a completely '
    'separate code path. After verification, the socket is tagged with user info '
    'for all subsequent real-time events (chat, WebRTC calls, group messaging).'
)

# --- Place 3 ---
doc.add_paragraph()
doc.add_heading('3.3 Frontend Sends Token on Every Request', level=2)
doc.add_paragraph('File: public/js/modules/api.js — HTTP requests')
code_block(
    'async get(endpoint) {\n'
    '    const response = await fetch(endpoint, {\n'
    '        headers: { \'Authorization\': `Bearer ${token}` }\n'
    '    });\n'
    '    ...\n'
    '}'
)
doc.add_paragraph('File: public/js/modules/socket.js — WebSocket authentication')
code_block(
    'connect() {\n'
    '    const token = localStorage.getItem(\'token\');\n'
    '    socket = io({ auth: { token } });\n'
    '    socket.on(\'connect\', () => {\n'
    '        socket.emit(\'authenticate\', token);\n'
    '    });\n'
    '}'
)


doc.add_page_break()


# ═══════════════════════════════════════
# SECTION 4: Full Token Lifecycle Diagram
# ═══════════════════════════════════════
doc.add_heading('4. Full Token Lifecycle', level=1)

lifecycle = """
 User types password
         │
         ▼
   POST /api/login  ──▶  authController.js  ──▶  jwt.sign({ id, username, isAdmin })
         │                                                     │
         ▼                                                     ▼
   Browser receives token                               Token saved in:
         │                                               • HTTP cookie (server-set)
         │                                               • localStorage (client-set)
         │
         ├──▶ Every HTTP request:
         │        Authorization: Bearer <token>
         │              │
         │              ▼
         │        middleware/auth.js
         │              │
         │              ▼
         │        jwt.verify(token, JWT_SECRET)
         │              │
         │              ▼
         │        req.user = { id, username, isAdmin }
         │              │
         │              ▼
         │        Controller processes request
         │
         │
         └──▶ Socket.IO connect:
                  socket.emit('authenticate', token)
                        │
                        ▼
                  socketManager.js
                        │
                        ▼
                  jwt.verify(token, JWT_SECRET)
                        │
                        ▼
                  socket.userId = id
                  socket.username = username
                        │
                        ▼
                  Real-time events (chat, calls, groups)
"""

p = doc.add_paragraph()
r = p.add_run(lifecycle.strip())
r.font.name = 'Consolas'
r.font.size = Pt(8.5)


doc.add_page_break()


# ═══════════════════════════════════════
# SECTION 5: Token Logout / Invalidation
# ═══════════════════════════════════════
doc.add_heading('5. Token Logout & Invalidation', level=1)

doc.add_paragraph('File: src/controllers/authController.js → logout()')
code_block(
    'exports.logout = asyncHandler(async (req, res) => {\n'
    '    await User.update({ is_logged_in: 0 }, { where: { id: req.user.id } });\n'
    '    res.clearCookie(\'token\');\n'
    '    res.json({ message: \'Logged out successfully\' });\n'
    '});'
)

doc.add_paragraph('File: public/js/auth.js → logout()')
code_block(
    'async function logout() {\n'
    '    await fetch(\'/api/auth/logout\', {\n'
    '        method: \'POST\',\n'
    '        headers: { \'Authorization\': `Bearer ${token}` }\n'
    '    });\n'
    '    localStorage.removeItem(\'token\');\n'
    '    localStorage.removeItem(\'userId\');\n'
    '    localStorage.removeItem(\'username\');\n'
    '    window.location.href = \'/\';\n'
    '}'
)

doc.add_paragraph('On logout, three things happen:')
styled_table(
    ['Action', 'Where', 'Effect'],
    [
        ['User.update({ is_logged_in: 0 })', 'Server (authController)', 'Marks user as logged out in DB'],
        ['res.clearCookie(\'token\')', 'Server → Browser', 'Deletes the HTTP-only cookie'],
        ['localStorage.removeItem(\'token\')', 'Browser (auth.js)', 'Removes token from localStorage'],
    ],
    'b45309'
)


doc.add_paragraph()
doc.add_paragraph()
doc.add_heading('6. Summary', level=1)

styled_table(
    ['Aspect', 'Details'],
    [
        ['Created in', 'authController.js (register + login)'],
        ['Signing secret', 'JWT_SECRET from .env via src/config/env.js'],
        ['Payload', '{ id, username, isAdmin }'],
        ['Expires', '24 hours'],
        ['Stored in', 'HTTP cookie (server-set) + localStorage (client-set)'],
        ['Verified for HTTP', 'src/middleware/auth.js → authenticateToken()'],
        ['Verified for WebSocket', 'src/sockets/socketManager.js → jwt.verify()'],
        ['Sent by frontend HTTP', 'public/js/modules/api.js → Authorization: Bearer header'],
        ['Sent by frontend WS', 'public/js/modules/socket.js → socket.emit(\'authenticate\')'],
        ['Invalidated on logout', 'Cookie cleared + localStorage removed + DB flag reset'],
    ],
    '2d3748'
)


# ── Save ──
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'SecureConnect_JWT_Analysis.docx')
doc.save(output_path)
print(f'Done! Saved to: {output_path}')
