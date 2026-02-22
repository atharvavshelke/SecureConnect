"""Generates a Word document explaining encryption/decryption in SecureConnect."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2); s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for lvl in range(1, 4):
    doc.styles[f'Heading {lvl}'].font.name = 'Calibri'
    doc.styles[f'Heading {lvl}'].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

def shade(cell, color):
    tc = cell._element.get_or_add_tcPr()
    tc.append(tc.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'}))

def styled_table(headers, rows, color='1e40af'):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.font.bold = True; r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); r.font.size = Pt(10)
        shade(c, color)
    for rd in rows:
        row = t.add_row()
        for i, txt in enumerate(rd):
            row.cells[i].text = txt
            for p in row.cells[i].paragraphs:
                for r in p.runs: r.font.size = Pt(10)
    return t

def code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    p.paragraph_format.left_indent = Cm(1)
    return p

def mono_diagram(text, size=8):
    p = doc.add_paragraph()
    r = p.add_run(text.strip())
    r.font.name = 'Consolas'
    r.font.size = Pt(size)
    return p

def bold_run(paragraph, text, color=None):
    r = paragraph.add_run(text)
    r.bold = True
    if color: r.font.color.rgb = color
    return r

def normal_run(paragraph, text):
    return paragraph.add_run(text)


# ═══════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
t = doc.add_heading('SecureConnect', level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t2 = doc.add_heading('Encryption & Decryption Analysis', level=1)
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('End-to-End Encryption: How Messages Are Secured')
r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x66,0x66,0x66)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('File: public/js/crypto.js (SecureEncryption class)')
r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x88,0x88,0x88)

doc.add_page_break()


# ═══════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc = [
    '1. Overview — What Algorithms Are Used',
    '2. Key Generation (Registration)',
    '   2.1 RSA Key Pair Generation',
    '   2.2 Private Key Protection with Password (PBKDF2 + AES-GCM)',
    '3. Message Encryption (Sending a Message)',
    '   3.1 Step-by-Step Encryption Flow',
    '   3.2 The Encrypted Package Format',
    '4. Message Decryption (Receiving a Message)',
    '   4.1 Step-by-Step Decryption Flow',
    '   4.2 Reading Your Own Sent Messages',
    '5. Group Chat Encryption',
    '   5.1 Group Key Distribution',
    '   5.2 Group Message Encryption/Decryption',
    '6. Key Storage & Sync Across Devices',
    '7. Complete System Diagram',
    '8. Algorithm Summary',
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()


# ═══════════════════════════════════════
# SECTION 1: Overview
# ═══════════════════════════════════════
doc.add_heading('1. Overview — What Algorithms Are Used', level=1)

doc.add_paragraph(
    'SecureConnect uses a hybrid encryption scheme combining asymmetric (RSA) and '
    'symmetric (AES) encryption. This is the same approach used by Signal, WhatsApp, '
    'and other E2E encrypted messengers.'
)

styled_table(
    ['Algorithm', 'Type', 'Key Size', 'Purpose', 'Where Used'],
    [
        ['RSA-OAEP', 'Asymmetric (Public/Private)', '2048-bit', 'Encrypt the AES key for each recipient', 'Key exchange between users'],
        ['AES-GCM', 'Symmetric', '256-bit', 'Encrypt the actual message content', 'Every message (1:1 and group)'],
        ['PBKDF2', 'Key Derivation', '600,000 iterations', 'Derive encryption key from password', 'Protecting private key for server storage'],
        ['SHA-256', 'Hash', '256-bit', 'Hash function for RSA-OAEP and PBKDF2', 'Internal to RSA and key derivation'],
    ],
    '2d3748'
)

doc.add_paragraph()
p = doc.add_paragraph()
bold_run(p, 'Why hybrid encryption? ', RGBColor(0xCC,0x00,0x00))
normal_run(p,
    'RSA can only encrypt small data (≤ 214 bytes with 2048-bit key). Messages can be much larger. '
    'So we encrypt the message with fast AES, then encrypt the small AES key with RSA. '
    'This gives us both speed (AES) and security (RSA key exchange).'
)

doc.add_paragraph()
p = doc.add_paragraph()
bold_run(p, 'E2E = End-to-End: ', RGBColor(0xCC,0x00,0x00))
normal_run(p,
    'ALL encryption/decryption happens in the BROWSER (public/js/crypto.js). '
    'The server NEVER sees plaintext messages. It only stores encrypted blobs.'
)


doc.add_page_break()


# ═══════════════════════════════════════
# SECTION 2: Key Generation
# ═══════════════════════════════════════
doc.add_heading('2. Key Generation (Registration)', level=1)

doc.add_heading('2.1 RSA Key Pair Generation', level=2)
doc.add_paragraph('File: public/js/crypto.js → generateKeyPair() (line 12-35)')
doc.add_paragraph('When a user registers, the browser generates an RSA-2048 key pair:')
code_block(
    'this.keyPair = await window.crypto.subtle.generateKey(\n'
    '    {\n'
    '        name: "RSA-OAEP",\n'
    '        modulusLength: 2048,\n'
    '        publicExponent: new Uint8Array([1, 0, 1]),  // 65537\n'
    '        hash: "SHA-256"\n'
    '    },\n'
    '    true,       // extractable\n'
    '    ["encrypt", "decrypt"]\n'
    ');'
)

mono_diagram("""
┌─────────────────────────────────────────────────────────────────────┐
│                   REGISTRATION - Key Generation                      │
│                                                                      │
│  Browser (crypto.js)                                                 │
│  ┌──────────────────────────────────────────────────────────┐       │
│  │  window.crypto.subtle.generateKey("RSA-OAEP", 2048-bit) │       │
│  │                     │                                     │       │
│  │                     ▼                                     │       │
│  │         ┌───────────────────────┐                        │       │
│  │         │   RSA Key Pair        │                        │       │
│  │         │  ┌─────────────────┐  │                        │       │
│  │         │  │  Public Key 🔓  │──┼────▶ Sent to Server    │       │
│  │         │  │  (can encrypt)  │  │      (stored in DB)    │       │
│  │         │  └─────────────────┘  │                        │       │
│  │         │  ┌─────────────────┐  │                        │       │
│  │         │  │  Private Key 🔑 │──┼────▶ Stored locally    │       │
│  │         │  │  (can decrypt)  │  │      (localStorage)    │       │
│  │         │  └─────────────────┘  │                        │       │
│  │         └───────────────────────┘                        │       │
│  └──────────────────────────────────────────────────────────┘       │
│                                                                      │
│  Server DB stores: public_key (Base64 string)                       │
│  Server NEVER has: raw private key                                   │
└─────────────────────────────────────────────────────────────────────┘
""")

doc.add_heading('2.2 Private Key Protection with Password (PBKDF2 + AES-GCM)', level=2)
doc.add_paragraph('File: public/js/crypto.js → encryptPrivateKeyWithPassword() (line 383-408)')
doc.add_paragraph(
    'To allow login from any device, the private key is encrypted with the user\'s password '
    'and sent to the server. The server stores the encrypted blob — it cannot decrypt it '
    'because it never knows the password in plaintext (bcrypt hash only).'
)

code_block(
    '// 1. Generate random salt (16 bytes)\n'
    'const salt = window.crypto.getRandomValues(new Uint8Array(16));\n\n'
    '// 2. Derive AES key from password using PBKDF2 (600,000 iterations!)\n'
    'const key = await this.deriveKeyFromPassword(password, salt);\n\n'
    '// 3. Generate random IV (12 bytes)\n'
    'const iv = window.crypto.getRandomValues(new Uint8Array(12));\n\n'
    '// 4. Export private key to raw bytes (PKCS8 format)\n'
    'const exportedPrivateKey = await crypto.subtle.exportKey("pkcs8", this.privateKey);\n\n'
    '// 5. Encrypt private key bytes with AES-GCM using the derived key\n'
    'const ciphertext = await crypto.subtle.encrypt(\n'
    '    { name: "AES-GCM", iv: iv },\n'
    '    key,\n'
    '    exportedPrivateKey\n'
    ');\n\n'
    '// 6. Bundle { salt, iv, ciphertext } as JSON → sent to server'
)

mono_diagram("""
┌─────────────────────────────────────────────────────────────────────┐
│              PRIVATE KEY PROTECTION FLOW                             │
│                                                                      │
│  User Password: "MyP@ss123"                                         │
│       │                                                              │
│       ▼                                                              │
│  ┌─────────────────────────────────────────────────┐                │
│  │  PBKDF2 (600,000 iterations, SHA-256)           │                │
│  │  Password + Random Salt (16 bytes)              │                │
│  │              │                                   │                │
│  │              ▼                                   │                │
│  │  Derived AES-256 Key                            │                │
│  └──────────────┬──────────────────────────────────┘                │
│                 │                                                    │
│                 ▼                                                    │
│  ┌─────────────────────────────────────────────────┐                │
│  │  AES-GCM Encrypt                                │                │
│  │  Key: derived from password                     │                │
│  │  IV:  random 12 bytes                           │                │
│  │  Data: raw private key (PKCS8)                  │                │
│  │              │                                   │                │
│  │              ▼                                   │                │
│  │  Encrypted Private Key Bundle                   │                │
│  │  { salt, iv, ciphertext }                       │                │
│  └──────────────┬──────────────────────────────────┘                │
│                 │                                                    │
│                 ▼                                                    │
│  POST /api/register → Server stores encrypted bundle in DB          │
│                                                                      │
│  ⚠️  Server CANNOT decrypt this — it only has bcrypt(password)      │
└─────────────────────────────────────────────────────────────────────┘
""")


doc.add_page_break()


# ═══════════════════════════════════════
# SECTION 3: Message Encryption
# ═══════════════════════════════════════
doc.add_heading('3. Message Encryption (Sending a Message)', level=1)

doc.add_heading('3.1 Step-by-Step Encryption Flow', level=2)
doc.add_paragraph('File: public/js/crypto.js → encryptMessage() (line 86-154)')
doc.add_paragraph(
    'When Alice sends a message to Bob, the following happens entirely in Alice\'s browser:'
)

styled_table(
    ['Step', 'Operation', 'Code Reference', 'Output'],
    [
        ['1', 'Generate a random AES-256 key (one-time use)', 'crypto.subtle.generateKey("AES-GCM", 256)', 'AES session key'],
        ['2', 'Generate a random IV (12 bytes)', 'crypto.getRandomValues(new Uint8Array(12))', 'Initialization Vector'],
        ['3', 'Encrypt message text with AES-GCM', 'crypto.subtle.encrypt("AES-GCM", aesKey, message)', 'ciphertext (encrypted message)'],
        ['4', 'Get Bob\'s public key from server', 'Fetched from /api/users endpoint', 'Bob\'s RSA public key'],
        ['5', 'Encrypt AES key with Bob\'s RSA public key', 'crypto.subtle.encrypt("RSA-OAEP", bobPubKey, aesKey)', 'encryptedKey (only Bob can decrypt)'],
        ['6', 'Encrypt AES key with Alice\'s OWN public key', 'crypto.subtle.encrypt("RSA-OAEP", alicePubKey, aesKey)', 'encryptedKeySelf (so Alice can read her own sent messages)'],
        ['7', 'Bundle everything as JSON', 'JSON.stringify({encryptedKey, encryptedKeySelf, iv, ciphertext})', 'Final encrypted package'],
    ],
    '1e40af'
)

doc.add_paragraph()

mono_diagram("""
┌─────────────────────────────────────────────────────────────────────────────┐
│                     MESSAGE ENCRYPTION (Alice → Bob)                        │
│                                                                             │
│  Alice's Browser                                                            │
│  ┌──────────────────────────────────────────────────────────────────────┐   │
│  │                                                                      │   │
│  │  "Hello Bob!" ─── plaintext message                                  │   │
│  │       │                                                              │   │
│  │       ▼                                                              │   │
│  │  ┌─────────────────────────────────────────────────────────────┐    │   │
│  │  │  Step 1: Generate random AES-256 key (one-time use)        │    │   │
│  │  │  Step 2: Generate random IV (12 bytes)                     │    │   │
│  │  └────────────────────┬────────────────────────────────────────┘    │   │
│  │                       │                                             │   │
│  │                       ▼                                             │   │
│  │  ┌─────────────────────────────────────────────────────────────┐    │   │
│  │  │  Step 3: AES-GCM Encrypt(message, aesKey, iv)              │    │   │
│  │  │          "Hello Bob!" ──▶ "x7Kp9mQ2..." (ciphertext)      │    │   │
│  │  └────────────────────┬────────────────────────────────────────┘    │   │
│  │                       │                                             │   │
│  │           ┌───────────┴───────────┐                                │   │
│  │           ▼                       ▼                                │   │
│  │  ┌────────────────────┐  ┌────────────────────┐                   │   │
│  │  │ Step 5: RSA-OAEP   │  │ Step 6: RSA-OAEP   │                   │   │
│  │  │ Encrypt AES key    │  │ Encrypt AES key    │                   │   │
│  │  │ with BOB's 🔓      │  │ with ALICE's 🔓    │                   │   │
│  │  │ public key         │  │ public key         │                   │   │
│  │  │     │              │  │     │              │                   │   │
│  │  │     ▼              │  │     ▼              │                   │   │
│  │  │ encryptedKey       │  │ encryptedKeySelf   │                   │   │
│  │  │ (only Bob's 🔑     │  │ (only Alice's 🔑   │                   │   │
│  │  │  can open this)    │  │  can open this)    │                   │   │
│  │  └────────┬───────────┘  └─────────┬──────────┘                   │   │
│  │           │                        │                               │   │
│  │           └─────────┬──────────────┘                               │   │
│  │                     ▼                                              │   │
│  │  ┌─────────────────────────────────────────────────────────────┐   │   │
│  │  │  Step 7: Build Encrypted Package (JSON)                    │   │   │
│  │  │  {                                                         │   │   │
│  │  │    "encryptedKey":     "Base64...",  ← for Bob              │   │   │
│  │  │    "encryptedKeySelf": "Base64...",  ← for Alice            │   │   │
│  │  │    "iv":               "Base64...",  ← random nonce         │   │   │
│  │  │    "ciphertext":       "Base64..."   ← encrypted message    │   │   │
│  │  │  }                                                         │   │   │
│  │  └────────────────────┬────────────────────────────────────────┘   │   │
│  └───────────────────────┼───────────────────────────────────────────┘   │
│                          │                                               │
│                          ▼                                               │
│  Socket.IO: emit('send-message', { toUserId, encryptedPackage })        │
│  Server stores encrypted blob in DB → NEVER sees plaintext              │
└─────────────────────────────────────────────────────────────────────────────┘
""")

doc.add_heading('3.2 The Encrypted Package Format', level=2)
doc.add_paragraph('This is what gets stored in the database for every message:')
code_block(
    '{\n'
    '    "encryptedKey":     "MIIBIjAN...",    // AES key encrypted with recipient\'s RSA pub key\n'
    '    "encryptedKeySelf": "QWxpY2Vm...",    // AES key encrypted with sender\'s RSA pub key\n'
    '    "iv":               "dG9rZW4t...",    // Random 12-byte initialization vector\n'
    '    "ciphertext":       "U2VjdXJl..."     // Message encrypted with AES-GCM\n'
    '}'
)

doc.add_page_break()


# ═══════════════════════════════════════
# SECTION 4: Message Decryption
# ═══════════════════════════════════════
doc.add_heading('4. Message Decryption (Receiving a Message)', level=1)

doc.add_heading('4.1 Step-by-Step Decryption Flow', level=2)
doc.add_paragraph('File: public/js/crypto.js → decryptMessage() (line 239-304)')
doc.add_paragraph(
    'When Bob receives the encrypted package, his browser decrypts it:'
)

styled_table(
    ['Step', 'Operation', 'What Happens'],
    [
        ['1', 'Parse JSON package', 'Extract encryptedKey, iv, ciphertext from the JSON blob'],
        ['2', 'RSA-OAEP Decrypt the AES key', 'Use Bob\'s PRIVATE key to decrypt encryptedKey → recovers the AES session key'],
        ['3', 'Import AES key', 'Import the raw AES key bytes as a CryptoKey object'],
        ['4', 'AES-GCM Decrypt the ciphertext', 'Use the AES key + IV to decrypt the ciphertext → recovers "Hello Bob!"'],
    ],
    '15803d'
)

mono_diagram("""
┌──────────────────────────────────────────────────────────────────────────┐
│                     MESSAGE DECRYPTION (Bob receives)                     │
│                                                                          │
│  Bob's Browser                                                           │
│  ┌───────────────────────────────────────────────────────────────────┐  │
│  │                                                                   │  │
│  │  Encrypted Package from server:                                   │  │
│  │  { encryptedKey, encryptedKeySelf, iv, ciphertext }               │  │
│  │       │                                                           │  │
│  │       ▼                                                           │  │
│  │  ┌─────────────────────────────────────────────────────────────┐  │  │
│  │  │  Step 2: RSA-OAEP Decrypt                                  │  │  │
│  │  │  encryptedKey + Bob's Private Key 🔑                       │  │  │
│  │  │           │                                                 │  │  │
│  │  │           ▼                                                 │  │  │
│  │  │  Recovered AES-256 session key                             │  │  │
│  │  └───────────┬─────────────────────────────────────────────────┘  │  │
│  │              │                                                    │  │
│  │              ▼                                                    │  │
│  │  ┌─────────────────────────────────────────────────────────────┐  │  │
│  │  │  Step 4: AES-GCM Decrypt                                   │  │  │
│  │  │  ciphertext + AES key + IV                                 │  │  │
│  │  │           │                                                 │  │  │
│  │  │           ▼                                                 │  │  │
│  │  │  "Hello Bob!"  ← Original plaintext message!               │  │  │
│  │  └─────────────────────────────────────────────────────────────┘  │  │
│  └───────────────────────────────────────────────────────────────────┘  │
└──────────────────────────────────────────────────────────────────────────┘
""")

doc.add_heading('4.2 Reading Your Own Sent Messages', level=2)
doc.add_paragraph(
    'When Alice looks at her own sent messages later, she cannot use encryptedKey '
    '(that was encrypted with Bob\'s public key). Instead, the code tries encryptedKeySelf '
    '(encrypted with Alice\'s public key). This is the fallback mechanism in decryptMessage():'
)
code_block(
    'try {\n'
    '    // Try recipient key first (works for Bob)\n'
    '    aesKeyBuffer = await crypto.subtle.decrypt("RSA-OAEP", privateKey, encryptedKey);\n'
    '} catch (recipientError) {\n'
    '    // Falls back to self-encrypted key (works for Alice reading her own sent msgs)\n'
    '    aesKeyBuffer = await crypto.subtle.decrypt("RSA-OAEP", privateKey, encryptedKeySelf);\n'
    '}'
)

mono_diagram("""
┌──────────────────────────────────────────────────────────────────────────┐
│              WHO CAN DECRYPT WHAT?                                        │
│                                                                          │
│  Encrypted Package for Alice→Bob message:                                │
│                                                                          │
│  encryptedKey     ─── encrypted with Bob's public key                    │
│       │                   │                                              │
│       │              Only Bob's private key 🔑 can decrypt this          │
│       │                                                                  │
│  encryptedKeySelf ─── encrypted with Alice's public key                  │
│       │                   │                                              │
│       │              Only Alice's private key 🔑 can decrypt this        │
│       │                                                                  │
│  Both unlock THE SAME AES key → which decrypts the ciphertext           │
│                                                                          │
│  ✅ Bob can read   → uses encryptedKey                                   │
│  ✅ Alice can read  → uses encryptedKeySelf (her own sent messages)      │
│  ❌ Server can't read → has no private keys                              │
│  ❌ Anyone else     → no matching private key                            │
└──────────────────────────────────────────────────────────────────────────┘
""")

doc.add_page_break()


# ═══════════════════════════════════════
# SECTION 5: Group Chat Encryption
# ═══════════════════════════════════════
doc.add_heading('5. Group Chat Encryption', level=1)

doc.add_heading('5.1 Group Key Distribution', level=2)
doc.add_paragraph(
    'Groups use a shared AES key. When a group is created, a random AES-256 key is '
    'generated and encrypted individually for each member using their RSA public key.'
)

doc.add_paragraph('File: public/js/crypto.js → generateGroupKey() (line 159-167)')
code_block(
    '// Generate shared AES key for the group\n'
    'const key = await crypto.subtle.generateKey(\n'
    '    { name: "AES-GCM", length: 256 },\n'
    '    true, ["encrypt", "decrypt"]\n'
    ');\n'
    'const exported = await crypto.subtle.exportKey("raw", key);\n'
    'return exported; // ArrayBuffer'
)

doc.add_paragraph('File: public/js/crypto.js → encryptGroupKeyForMember() (line 170-178)')
code_block(
    '// Encrypt group key for each member individually\n'
    'async encryptGroupKeyForMember(groupKeyBuffer, memberPublicKeyString) {\n'
    '    const publicKey = await this.importPublicKey(memberPublicKeyString);\n'
    '    const encrypted = await crypto.subtle.encrypt(\n'
    '        { name: "RSA-OAEP" }, publicKey, groupKeyBuffer\n'
    '    );\n'
    '    return this.arrayBufferToBase64(encrypted);\n'
    '}'
)

mono_diagram("""
┌──────────────────────────────────────────────────────────────────────────┐
│                   GROUP KEY DISTRIBUTION                                  │
│                                                                          │
│  Group Creator (Alice) generates a random AES-256 group key              │
│       │                                                                  │
│       │         For each member:                                         │
│       ├─────▶ RSA-Encrypt(groupKey, Alice's pubKey)  → Alice's copy      │
│       ├─────▶ RSA-Encrypt(groupKey, Bob's pubKey)    → Bob's copy        │
│       └─────▶ RSA-Encrypt(groupKey, Carol's pubKey)  → Carol's copy      │
│                                                                          │
│  Each member gets the SAME group key, but encrypted with THEIR pubkey    │
│  Only THEY can decrypt it with their private key                         │
│                                                                          │
│  Server stores: { memberId: encryptedGroupKey } per member               │
│  Server CANNOT decrypt any of them                                       │
└──────────────────────────────────────────────────────────────────────────┘
""")


doc.add_heading('5.2 Group Message Encryption/Decryption', level=2)
doc.add_paragraph('Once a member has the group key, encryption is simpler than 1:1 chat:')

styled_table(
    ['Operation', 'Function', 'How It Works'],
    [
        ['Encrypt', 'encryptWithGroupKey(message, groupKeyBuffer)', '1. Import group AES key\n2. Generate random IV\n3. AES-GCM encrypt message\n4. Return { iv, ciphertext, isGroup: true }'],
        ['Decrypt', 'decryptWithGroupKey(encrypted, groupKeyBuffer)', '1. Parse { iv, ciphertext } from JSON\n2. Import group AES key\n3. AES-GCM decrypt → plaintext'],
    ],
    '6d28d9'
)

doc.add_paragraph()
p = doc.add_paragraph()
bold_run(p, 'Key difference from 1:1: ', RGBColor(0xCC,0x00,0x00))
normal_run(p,
    'In 1:1 chat, a NEW AES key is generated per message. In group chat, ALL members '
    'share the SAME AES key, so no RSA step is needed per message — just AES encrypt/decrypt.'
)


doc.add_page_break()


# ═══════════════════════════════════════
# SECTION 6: Key Storage & Sync
# ═══════════════════════════════════════
doc.add_heading('6. Key Storage & Sync Across Devices', level=1)

styled_table(
    ['Key', 'Where Stored', 'Format', 'Who Can Access It'],
    [
        ['Public Key', 'Server DB (users.public_key)', 'Base64 SPKI', 'Anyone — it\'s meant to be public'],
        ['Private Key (raw)', 'Browser localStorage', 'Base64 PKCS8', 'Only the user on that device'],
        ['Private Key (encrypted)', 'Server DB (users.encrypted_private_key)', 'JSON { salt, iv, ciphertext }', 'Only the user who knows the password'],
    ],
    'b45309'
)

doc.add_paragraph()
doc.add_paragraph('When a user logs in on a NEW device:')

mono_diagram("""
┌──────────────────────────────────────────────────────────────────────────┐
│                   LOGIN ON NEW DEVICE                                     │
│                                                                          │
│  1. User enters password → POST /api/login                               │
│  2. Server returns: { ..., encryptedPrivateKey: { salt, iv, ciphertext } │
│  3. Browser runs decryptPrivateKeyWithPassword(password, bundle):        │
│                                                                          │
│     password + salt                                                      │
│         │                                                                │
│         ▼                                                                │
│     PBKDF2 (600,000 iterations)                                         │
│         │                                                                │
│         ▼                                                                │
│     Derived AES key                                                      │
│         │                                                                │
│         ▼                                                                │
│     AES-GCM Decrypt(ciphertext, derivedKey, iv)                          │
│         │                                                                │
│         ▼                                                                │
│     Raw private key (PKCS8) → imported as CryptoKey                      │
│         │                                                                │
│         ▼                                                                │
│     Stored in localStorage for future use                                │
│                                                                          │
│  ✅ User can now decrypt all historical messages on any device           │
└──────────────────────────────────────────────────────────────────────────┘
""")

doc.add_page_break()


# ═══════════════════════════════════════
# SECTION 7: Complete System Diagram
# ═══════════════════════════════════════
doc.add_heading('7. Complete System Diagram', level=1)
doc.add_paragraph('This shows the full encryption lifecycle from registration to message exchange:')

mono_diagram("""
═══════════════════════════════════════════════════════════════════════════════════════
                      COMPLETE E2E ENCRYPTION SYSTEM
═══════════════════════════════════════════════════════════════════════════════════════

  REGISTRATION                                    LOGIN (new device)
  ────────────                                    ─────────────────
  Browser:                                        Browser:
  1. Generate RSA-2048 key pair                   1. Receive encrypted_private_key from server
  2. Export public key → send to server           2. PBKDF2(password, salt) → derived AES key
  3. PBKDF2(password, salt) → derived key         3. AES-GCM decrypt → recover private key
  4. AES-GCM encrypt(privateKey, derivedKey)      4. Store in localStorage
  5. Send encrypted private key to server         5. Ready to decrypt messages!
  6. Store raw private key in localStorage

═══════════════════════════════════════════════════════════════════════════════════════

  1:1 MESSAGE (Alice → Bob)
  ─────────────────────────
  Alice's Browser:                               Bob's Browser:
  1. Generate random AES-256 key                 1. Receive encrypted package
  2. Generate random IV (12 bytes)               2. RSA decrypt encryptedKey
  3. AES-GCM encrypt(message, aesKey, iv)           using Bob's private key
  4. RSA encrypt(aesKey, Bob's pubKey)              → recover AES key
     → encryptedKey                              3. AES-GCM decrypt(ciphertext,
  5. RSA encrypt(aesKey, Alice's pubKey)             aesKey, iv)
     → encryptedKeySelf                             → "Hello Bob!"
  6. Send { encryptedKey, encryptedKeySelf,
            iv, ciphertext } via Socket.IO

  Server: Stores encrypted blob. CANNOT read it.

═══════════════════════════════════════════════════════════════════════════════════════

  GROUP MESSAGE
  ─────────────
  Group Creation:                                Group Messaging:
  1. Generate random AES-256 group key            1. Sender: AES-GCM encrypt(message,
  2. For each member:                                groupKey, randomIV)
     RSA encrypt(groupKey, member's pubKey)       2. Send { iv, ciphertext, isGroup }
  3. Store encrypted copies on server             3. Receiver: decrypt own copy of
                                                     groupKey → AES-GCM decrypt → text

═══════════════════════════════════════════════════════════════════════════════════════
""", 7.5)


doc.add_page_break()


# ═══════════════════════════════════════
# SECTION 8: Algorithm Summary
# ═══════════════════════════════════════
doc.add_heading('8. Algorithm Summary', level=1)

styled_table(
    ['Component', 'Algorithm', 'Key Size', 'Purpose', 'Where in Code'],
    [
        ['Key Pair', 'RSA-OAEP + SHA-256', '2048-bit', 'Asymmetric encryption for key exchange', 'generateKeyPair() line 12'],
        ['Message Body', 'AES-GCM', '256-bit', 'Fast symmetric encryption of message text', 'encryptMessage() line 86'],
        ['Message Key Exchange', 'RSA-OAEP', '2048-bit', 'Encrypt per-message AES key for recipient', 'encryptMessage() line 119'],
        ['Self-Read Key', 'RSA-OAEP', '2048-bit', 'Encrypt AES key for sender (read own msgs)', 'encryptMessage() line 131'],
        ['Private Key Protection', 'PBKDF2 → AES-GCM', '256-bit derived', 'Password-encrypt private key for server storage', 'encryptPrivateKeyWithPassword() line 383'],
        ['Password Derivation', 'PBKDF2', '600K iterations', 'Slow hash to resist brute-force attacks', 'deriveKeyFromPassword() line 358'],
        ['Group Key', 'AES-GCM', '256-bit', 'Shared symmetric key for group messages', 'generateGroupKey() line 159'],
        ['Group Key Distribution', 'RSA-OAEP', '2048-bit', 'Encrypt group key per member', 'encryptGroupKeyForMember() line 170'],
        ['IV/Nonce', 'Random', '96-bit (12 bytes)', 'Unique per encryption to prevent pattern attacks', 'crypto.getRandomValues()'],
        ['Salt', 'Random', '128-bit (16 bytes)', 'Unique per password derivation', 'crypto.getRandomValues()'],
    ],
    '2d3748'
)

doc.add_paragraph()
p = doc.add_paragraph()
bold_run(p, 'Security Guarantee: ', RGBColor(0xCC,0x00,0x00))
normal_run(p,
    'All encryption and decryption runs in the browser using the Web Crypto API. '
    'The server NEVER sees plaintext messages, raw private keys, or user passwords '
    '(only bcrypt hashes). Even if the server database is compromised, messages '
    'remain encrypted and unreadable without the users\' private keys.'
)


# ── Save ──
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'SecureConnect_Encryption_Analysis.docx')
doc.save(output_path)
print(f'Done! Saved to: {output_path}')
